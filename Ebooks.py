import requests
from bs4 import BeautifulSoup
import random
import logging
import matplotlib.pyplot as plt
from docx import Document

class Ebooks:
    def __init__(self, app_instance):
        self.book_title = None
        self.author_name = None
        self.first_chapter = None
        self.paragraph_lengths = None
        self.app = app_instance 

        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    # Method to download a book from the Internet
    def download_ebook(self):
        try:
            url = "http://self.gutenberg.org"
            response = requests.get(url)
            if response.status_code == 200:
                html = response.text
                soup = BeautifulSoup(html, 'html.parser')

                ebook_links = soup.find_all('a', href=True, text='Plain Text UTF-8')
                if ebook_links:
                    random_link = random.choice(ebook_links)
                    ebook_url = random_link['href']

                    ebook_response = requests.get(ebook_url)
                    if ebook_response.status_code == 200:
                        ebook_html = ebook_response.text
                        self.extract_book_info(ebook_html)
                        self.logger.info("Ebook downloaded successfully.")
                        self.app.append_log("Ebook downloaded successfully.")
                    else:
                        self.logger.error(f"Failed to download ebook from {ebook_url}.")
                        self.app.append_log(f"Failed to download ebook from {ebook_url}.")
                else:
                    self.logger.warning("No ebooks found on the page.")
                    self.app.append_log("No ebooks found on the page.")
            else:
                self.logger.error(f"Failed to access {url}. Status code: {response.status_code}")
                self.app.append_log(f"Failed to access {url}. Status code: {response.status_code}")
        except Exception as e:
            self.logger.error(f"An error occurred: {str(e)}")
            self.app.append_log(f"An error occurred: {str(e)}")

    # Method to extract the book title, author name and first chapter from the HTML
    def extract_book_info(self, html):
        soup = BeautifulSoup(html, 'html.parser')
        self.book_title = soup.find('h1').text
        self.author_name = soup.find('h2').text
        self.first_chapter = soup.find('p').text

    # Method to count the number of words in each paragraph of the first chapter
    def count_paragraph_words(self):
        paragraphs = self.first_chapter.split('\n\n')
        self.paragraph_lengths = [self.round_word_count(len(p.split())) for p in paragraphs]

    # Method to round the word count to the nearest ten
    def round_word_count(self, count):
        return round(count, -1)

    # Method to plot the distribution of paragraph lengths
    def plot_paragraph_lengths(self):
        plt.hist(self.paragraph_lengths, bins=len(set(self.paragraph_lengths)))
        plt.xlabel('Nombre de mots par paragraphe')
        plt.ylabel('Nombre de paragraphes')
        plt.title('Distribution des longueurs des paragraphes')
        plt.savefig('distribution_paragraphes.png')
        plt.close()

    # Method to download an image from the Internet
    def download_image(self, url, filename):
        response = requests.get(url)
        with open(filename, 'wb') as f:
            f.write(response.content)

    # Method to resize an image
    def resize_image(self, image_path, output_path, width, height):
        pass

    # Method to rotate and paste an image
    def rotate_and_paste_image(self, image_path1, image_path2, output_path, angle):
        pass

    # Method to create a Word document
    def create_word_document(self, filename):
        doc = Document()
        doc.add_heading('Titre du livre : ' + self.book_title, level=1)
        doc.add_paragraph('Auteur du livre : ' + self.author_name)
        doc.add_picture('image1.jpg', width=Document().page_width)
        doc.add_heading('Auteur du rapport : Votre nom', level=2)
        doc.add_page_break()

        doc.add_heading('Distribution des longueurs des paragraphes', level=1)
        doc.add_paragraph('Description de l\'image : explication de l\'intrigue...')
        self.plot_paragraph_lengths()
        doc.add_picture('distribution_paragraphes.png', width=Document().page_width)
        doc.save(filename)
